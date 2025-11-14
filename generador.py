from docx import Document
import os

def generar_desistimiento(datos):
    """Sustituye variables en el formato de desistimiento"""
    
    # Ruta al documento plantilla
    template_path = "PactoMaker/formatos/Formato DESISTIMIENTO.docx"
    
    # Carpeta donde guardar el resultado
    output_dir = "cumplimientos_generados"
    
    # Crear carpeta si no existe
    os.makedirs(output_dir, exist_ok=True)
    
    # Nombre del archivo de salida
    expediente = datos.get('NUMERO_EXPEDIENTE', 'sin_numero').replace('/', '_')
    output_path = f"{output_dir}/Desistimiento_{expediente}.docx"
    
    print(f"📄 Leyendo plantilla: {template_path}")
    
    # Abrir el documento
    doc = Document(template_path)
    
    print("🔍 Buscando variables para sustituir...")
    
    # Sustituir en párrafos
    for paragraph in doc.paragraphs:
        for key, value in datos.items():
            if f'{{{key}}}' in paragraph.text:
                print(f"   Sustituyendo {{{key}}} → {value}")
                paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))
    
    # Sustituir en tablas (si el documento tiene tablas)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in datos.items():
                    if f'{{{key}}}' in cell.text:
                        print(f"   Sustituyendo en tabla: {{{key}}} → {value}")
                        cell.text = cell.text.replace(f'{{{key}}}', str(value))
    
    # Guardar el documento
    doc.save(output_path)
    print(f"✅ Documento generado: {output_path}")
    return output_path

# Esto es solo para probar - lo quitaremos después
if __name__ == "__main__":
    print("🧪 PROBANDO GENERADOR...")
    
    datos_prueba = {
    'NUMERO_EXPEDIENTE': '123/2024',
    'TIPO_PROCEDIMIENTO': 'LABORAL',
    'FECHA_DE_ACUERDO': '28 de octubre de 2024',
    'HORA_DE_COMPARECENCIA': '10:30 horas',
    'PARTE_ACTORA': 'JUAN PÉREZ',
    'PARTE_DEMANDADA': 'EMPRESA XYZ',
    'AUDIENCIA': 'AUDIENCIA PRELIMINAR',
    'FECHA_AUDIENCIA': '15 de noviembre de 2024',
    'SECRETARIO_DE_ACUERDOS': 'LIC. MARÍA GÓMEZ',
    'FECHA_DE_PUBLICACION': '29 de octubre de 2024'
}
    
    generar_desistimiento(datos_prueba)