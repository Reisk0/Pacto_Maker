from django.shortcuts import render
from django.http import HttpResponse, FileResponse
import os
from docx import Document
from datetime import datetime, timedelta

def formatear_fecha(fecha_str):
    """Convertir fecha YYYY-MM-DD a texto en español completo"""
    if not fecha_str:
        return ""
    
    try:
        # Diccionario de meses en español
        meses = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        
        # Diccionario de números a texto
        numeros = {
            1: 'uno', 2: 'dos', 3: 'tres', 4: 'cuatro', 5: 'cinco',
            6: 'seis', 7: 'siete', 8: 'ocho', 9: 'nueve', 10: 'diez',
            11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
            16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve',
            20: 'veinte', 21: 'veintiuno', 22: 'veintidós', 23: 'veintitrés', 
            24: 'veinticuatro', 25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete',
            28: 'veintiocho', 29: 'veintinueve', 30: 'treinta', 31: 'treinta y uno'
        }
        
        fecha = datetime.strptime(fecha_str, '%Y-%m-%d')
        
        # Convertir día a texto
        dia_texto = numeros.get(fecha.day, str(fecha.day))
        
        # Convertir año a texto
        año = fecha.year
        if año == 2024:
            año_texto = "dos mil veinticuatro"
        elif año == 2025:
            año_texto = "dos mil veinticinco" 
        elif año == 2026:
            año_texto = "dos mil veintiséis"
        else:
            año_texto = str(año)  # Para otros años usar número
        
        return f"{dia_texto} de {meses[fecha.month]} de {año_texto}"
    
    except Exception as e:
        return fecha_str
    """Convertir fecha YYYY-MM-DD a texto en español completo"""
    if not fecha_str:
        return ""
    
    try:
        meses = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        
        fecha = datetime.strptime(fecha_str, '%Y-%m-%d')
        return f"{fecha.day} de {meses[fecha.month]} de {fecha.year}"
    
    except Exception as e:
        return fecha_str

def formatear_fecha_publicacion(fecha_acuerdo_str):
    """Calcular fecha de publicación (3 días después de la fecha de acuerdo)"""
    if not fecha_acuerdo_str:
        return ""
    
    try:
        fecha_acuerdo = datetime.strptime(fecha_acuerdo_str, '%Y-%m-%d')
        fecha_publicacion = fecha_acuerdo + timedelta(days=3)
        
        meses = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        
        numeros = {
            1: 'uno', 2: 'dos', 3: 'tres', 4: 'cuatro', 5: 'cinco',
            6: 'seis', 7: 'siete', 8: 'ocho', 9: 'nueve', 10: 'diez',
            11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
            16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve',
            20: 'veinte', 21: 'veintiuno', 22: 'veintidós', 23: 'veintitrés', 
            24: 'veinticuatro', 25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete',
            28: 'veintiocho', 29: 'veintinueve', 30: 'treinta', 31: 'treinta y uno'
        }
        
        dia_texto = numeros.get(fecha_publicacion.day, str(fecha_publicacion.day))
        
        año = fecha_publicacion.year
        if año == 2024:
            año_texto = "dos mil veinticuatro"
        elif año == 2025:
            año_texto = "dos mil veinticinco"
        elif año == 2026:
            año_texto = "dos mil veintiséis"
        else:
            año_texto = str(año)
        
        return f"{dia_texto} de {meses[fecha_publicacion.month]} de {año_texto}"
    
    except Exception as e:
        return ""
    """Calcular fecha de publicación (3 días después de la fecha de acuerdo)"""
    if not fecha_acuerdo_str:
        return ""
    
    try:
        fecha_acuerdo = datetime.strptime(fecha_acuerdo_str, '%Y-%m-%d')
        fecha_publicacion = fecha_acuerdo + timedelta(days=3)
        
        meses = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        
        return f"{fecha_publicacion.day} de {meses[fecha_publicacion.month]} de {fecha_publicacion.year}"
    
    except Exception as e:
        return ""

def formatear_hora(hora_str):
    """Convertir hora HH:MM a texto en español completo"""
    if not hora_str:
        return ""
    
    try:
        # Diccionario completo de números a texto
        numeros = {
            0: 'cero', 1: 'una', 2: 'dos', 3: 'tres', 4: 'cuatro', 5: 'cinco',
            6: 'seis', 7: 'siete', 8: 'ocho', 9: 'nueve', 10: 'diez',
            11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
            16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve',
            20: 'veinte', 21: 'veintiuna', 22: 'veintidós', 23: 'veintitrés', 
            24: 'veinticuatro', 25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete',
            28: 'veintiocho', 29: 'veintinueve', 30: 'treinta', 31: 'treinta y una',
            32: 'treinta y dos', 33: 'treinta y tres', 34: 'treinta y cuatro', 
            35: 'treinta y cinco', 36: 'treinta y seis', 37: 'treinta y siete',
            38: 'treinta y ocho', 39: 'treinta y nueve', 40: 'cuarenta',
            41: 'cuarenta y una', 42: 'cuarenta y dos', 43: 'cuarenta y tres',
            44: 'cuarenta y cuatro', 45: 'cuarenta y cinco', 46: 'cuarenta y seis',
            47: 'cuarenta y siete', 48: 'cuarenta y ocho', 49: 'cuarenta y nueve',
            50: 'cincuenta', 51: 'cincuenta y una', 52: 'cincuenta y dos',
            53: 'cincuenta y tres', 54: 'cincuenta y cuatro', 55: 'cincuenta y cinco',
            56: 'cincuenta y seis', 57: 'cincuenta y siete', 58: 'cincuenta y ocho',
            59: 'cincuenta y nueve'
        }
        
        hora = datetime.strptime(hora_str, '%H:%M')
        
        # Convertir a formato 12 horas
        hora_12 = hora.hour if hora.hour <= 12 else hora.hour - 12
        periodo = 'a.m.' if hora.hour < 12 else 'p.m.'
        
        # Casos especiales
        if hora.hour == 12 and hora.minute == 0:
            return "doce horas"
        if hora.hour == 0 and hora.minute == 0:
            return "doce horas"
        
        # Convertir a texto
        hora_texto = numeros.get(hora_12, str(hora_12))
        minuto_texto = numeros.get(hora.minute, str(hora.minute))
        
        # Formato: "siete horas con veintiséis minutos"
        if hora.minute == 0:
            return f"{hora_texto} horas"
        else:
            return f"{hora_texto} horas con {minuto_texto} minutos"
    
    except Exception as e:
        return hora_str
    """Convertir hora HH:MM a texto en español completo"""
    if not hora_str:
        return ""
    
    try:
        # Diccionario completo de números a texto
        numeros = {
            0: 'cero', 1: 'una', 2: 'dos', 3: 'tres', 4: 'cuatro', 5: 'cinco',
            6: 'seis', 7: 'siete', 8: 'ocho', 9: 'nueve', 10: 'diez',
            11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
            16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve',
            20: 'veinte', 21: 'veintiuna', 22: 'veintidós', 23: 'veintitrés', 
            24: 'veinticuatro', 25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete',
            28: 'veintiocho', 29: 'veintinueve', 30: 'treinta', 31: 'treinta y una',
            32: 'treinta y dos', 33: 'treinta y tres', 34: 'treinta y cuatro', 
            35: 'treinta y cinco', 36: 'treinta y seis', 37: 'treinta y siete',
            38: 'treinta y ocho', 39: 'treinta y nueve', 40: 'cuarenta',
            41: 'cuarenta y una', 42: 'cuarenta y dos', 43: 'cuarenta y tres',
            44: 'cuarenta y cuatro', 45: 'cuarenta y cinco', 46: 'cuarenta y seis',
            47: 'cuarenta y siete', 48: 'cuarenta y ocho', 49: 'cuarenta y nueve',
            50: 'cincuenta', 51: 'cincuenta y una', 52: 'cincuenta y dos',
            53: 'cincuenta y tres', 54: 'cincuenta y cuatro', 55: 'cincuenta y cinco',
            56: 'cincuenta y seis', 57: 'cincuenta y siete', 58: 'cincuenta y ocho',
            59: 'cincuenta y nueve'
        }
        
        hora = datetime.strptime(hora_str, '%H:%M')
        
        # Determinar si es "las" o "la"
        articulo = "las" if hora.hour != 1 else "la"
        
        # Convertir a formato 12 horas
        hora_12 = hora.hour if hora.hour <= 12 else hora.hour - 12
        periodo = 'a.m.' if hora.hour < 12 else 'p.m.'
        
        # Casos especiales
        if hora.hour == 12 and hora.minute == 0:
            return "las doce del día"
        if hora.hour == 0 and hora.minute == 0:
            return "las doce de la noche"
        
        # Convertir a texto
        hora_texto = numeros.get(hora_12, str(hora_12))
        minuto_texto = numeros.get(hora.minute, str(hora.minute))
        
        return f"{articulo} {hora_texto} y {minuto_texto} {periodo}"
    
    except Exception as e:
        return hora_str

def generar_documento_generico(tipo_documento, datos):
    """Función ROBUSTA para generar documentos - INCLUYE ENCABEZADOS"""
    
    plantillas = {
        'desistimiento': "PactoMaker/formatos/Formato DESISTIMIENTO.docx",
        'desistimiento_sin_efectos': "PactoMaker/formatos/Formato DESISTIMIENTO_SIN_EFECTOS_AUDIENCIA.docx",
        'pago_efectivo': "PactoMaker/formatos/Formato PAGO_EFECTIVO.docx",
        'pago_transferencia': "PactoMaker/formatos/Formato PAGO_TRANSFERENCIA.docx", 
        'pago_cheque': "PactoMaker/formatos/Formato PAGO_CHEQUE.docx",
        'pago_cumplido': "PactoMaker/formatos/Formato SE_DA_POR_PAGADO.docx",
        'pago_parcial_cheque': "PactoMaker/formatos/Formato PAGO_PARCIAL_CHEQUE.docx",
    }
    
    template_path = plantillas.get(tipo_documento)
    if not template_path or not os.path.exists(template_path):
        return f"Error: Plantilla para {tipo_documento} no encontrada"
    
    output_dir = "documentos_generados"
    os.makedirs(output_dir, exist_ok=True)
    
    expediente = datos.get('NUMERO_EXPEDIENTE', 'sin_numero').replace('/', '_')
    output_path = f"{output_dir}/{tipo_documento}_{expediente}.docx"
    
    try:
        doc = Document(template_path)
        
        # ✅ 1. REEMPLAZAR EN ENCABEZADOS
        for section in doc.sections:
            # Encabezado
            if section.header:
                for paragraph in section.header.paragraphs:
                    for key, value in datos.items():
                        placeholder = f'{{{key}}}'
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Pie de página  
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for key, value in datos.items():
                        placeholder = f'{{{key}}}'
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # ✅ 2. REEMPLAZO EN PÁRRAFOS PRINCIPALES
        for paragraph in doc.paragraphs:
            for key, value in datos.items():
                placeholder = f'{{{key}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # ✅ 3. REEMPLAZO EN TABLAS  
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in datos.items():
                        placeholder = f'{{{key}}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        return f"Error al procesar documento: {str(e)}"
    """Función ROBUSTA para generar documentos"""
    
    plantillas = {
        'desistimiento': "PactoMaker/formatos/Formato DESISTIMIENTO.docx",
        'desistimiento_sin_efectos': "PactoMaker/formatos/Formato DESISTIMIENTO_SIN_EFECTOS_AUDIENCIA.docx",
        'pago_efectivo': "PactoMaker/formatos/Formato PAGO_EFECTIVO.docx",
        'pago_transferencia': "PactoMaker/formatos/Formato PAGO_TRANSFERENCIA.docx", 
        'pago_cheque': "PactoMaker/formatos/Formato PAGO_CHEQUE.docx",
        'pago_cumplido': "PactoMaker/formatos/Formato SE_DA_POR_PAGADO.docx",
    }
    
    template_path = plantillas.get(tipo_documento)
    if not template_path or not os.path.exists(template_path):
        return f"Error: Plantilla para {tipo_documento} no encontrada"
    
    output_dir = "documentos_generados"
    os.makedirs(output_dir, exist_ok=True)
    
    expediente = datos.get('NUMERO_EXPEDIENTE', 'sin_numero').replace('/', '_')
    output_path = f"{output_dir}/{tipo_documento}_{expediente}.docx"
    
    try:
        doc = Document(template_path)
        
        # REEMPLAZO ROBUSTO EN PÁRRAFOS
        for paragraph in doc.paragraphs:
            for key, value in datos.items():
                placeholder = f'{{{key}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # REEMPLAZO ROBUSTO EN TABLAS  
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in datos.items():
                        placeholder = f'{{{key}}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        return f"Error al procesar documento: {str(e)}"

def home(request):
    """Página principal con selector de tipo de documento"""
    return render(request, 'generador_docs/home.html')

def generar_documento(request):
    """Procesar el formulario y generar documento"""
    if request.method == 'POST':
        tipo_documento = request.POST.get('TIPO_DOCUMENTO', 'desistimiento')
        
        # Recoger y formatear datos
        datos = {
            'NUMERO_EXPEDIENTE': request.POST.get('NUMERO_EXPEDIENTE', ''),
            'TIPO_PROCEDIMIENTO': request.POST.get('TIPO_PROCEDIMIENTO', ''),
            'PARTE_ACTORA': request.POST.get('PARTE_ACTORA', ''),
            'PARTE_DEMANDADA': request.POST.get('PARTE_DEMANDADA', ''),
            'SECRETARIO_DE_ACUERDOS': request.POST.get('SECRETARIO_DE_ACUERDOS', ''),
            
            # FECHAS Y HORAS FORMATEADAS
            'FECHA_DE_ACUERDO': formatear_fecha(request.POST.get('FECHA_DE_ACUERDO', '')),
            'HORA_DE_COMPARECENCIA': formatear_hora(request.POST.get('HORA_DE_COMPARECENCIA', '')),
            'FECHA_DE_PUBLICACION': formatear_fecha_publicacion(request.POST.get('FECHA_DE_ACUERDO', '')),
            'FECHA_AUDIENCIA': formatear_fecha(request.POST.get('FECHA_AUDIENCIA', '')),
            'FECHA_CONVENIO': formatear_fecha(request.POST.get('FECHA_CONVENIO', '')),
            
            # CAMPOS DE AUDIENCIA
            'AUDIENCIA': request.POST.get('AUDIENCIA', ''),
            
            # CAMPOS PARA PAGOS
            'PERSONALIDAD': request.POST.get('PERSONALIDAD', ''),
            'REPRESENTANTE_DEMANDADA': request.POST.get('REPRESENTANTE_DEMANDADA', ''),
            'CANTIDAD': request.POST.get('CANTIDAD', ''),
            
            # NUEVOS CAMPOS
            'NOMBRE_COMPARECIENTE': request.POST.get('PARTE_ACTORA', ''),
            'DOCUMENTO_IDENTIFICACION': request.POST.get('DOCUMENTO_IDENTIFICACION', 'credencial para votar vigente'),
            'AUTORIDAD_EXPEDIDORA': request.POST.get('AUTORIDAD_EXPEDIDORA', 'Instituto Nacional Electoral'),
            'NOMBRE_JUEZ': request.POST.get('NOMBRE_JUEZ', 'Karla María Gallegos Castañeda'),
            'CARGO_JUEZ': request.POST.get('CARGO_JUEZ', 'Jueza Segunda adscrita al Tribunal Laboral en el Estado'),
            'NOMBRE_ASISTENTE': request.POST.get('NOMBRE_ASISTENTE', 'Mayra Fernanda Lujan Serna'),
            'CARGO_ASISTENTE': request.POST.get('CARGO_ASISTENTE', 'Secretaria de Acuerdos'),

            # ✅ NUEVO CAMPO PARA CHEQUE/TRANSFERENCIA
            'TITULO_CREDITO': request.POST.get('TITULO_CREDITO', ''),

                # ✅ NUEVO CAMPO PARA TRANSFERENCIA
            'CUENTA_BANCARIA': request.POST.get('CUENTA_BANCARIA', ''),
            'PAGO_PARCIAL': request.POST.get('PAGO_PARCIAL', ''),
            'CANTIDAD_PENDIENTE': request.POST.get('CANTIDAD_PENDIENTE', ''),
        }
        
        try:
            archivo_generado = generar_documento_generico(tipo_documento, datos)
            
            if archivo_generado.startswith("Error:"):
                return HttpResponse(f"❌ {archivo_generado}")
            else:
                if os.path.exists(archivo_generado):
                    response = FileResponse(
                        open(archivo_generado, 'rb'),
                        as_attachment=True,
                        filename=os.path.basename(archivo_generado)
                    )
                    return response
                else:
                    return HttpResponse(f"✅ Documento generado: {archivo_generado}")
        
        except Exception as e:
            return HttpResponse(f"❌ Error al generar documento: {str(e)}")
    
    return HttpResponse("Método no permitido")