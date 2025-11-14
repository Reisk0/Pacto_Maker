from docx import Document

def encontrar_variables(documento_path):
    """Encuentra todas las variables {VARIABLE} en un documento DOCX"""
    doc = Document(documento_path)
    variables_encontradas = set()
    
    print("🔍 Buscando variables en el documento...")
    
    # Buscar en párrafos
    for paragraph in doc.paragraphs:
        texto = paragraph.text
        inicio = texto.find('{')
        while inicio != -1:
            fin = texto.find('}', inicio)
            if fin != -1:
                variable = texto[inicio+1:fin]
                variables_encontradas.add(variable)
                print(f"   Encontrada: {{{variable}}}")
            inicio = texto.find('{', fin)
    
    # Buscar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto = cell.text
                inicio = texto.find('{')
                while inicio != -1:
                    fin = texto.find('}', inicio)
                    if fin != -1:
                        variable = texto[inicio+1:fin]
                        variables_encontradas.add(variable)
                        print(f"   Encontrada en tabla: {{{variable}}}")
                    inicio = texto.find('{', fin)
    
    print(f"\n📋 TOTAL de variables encontradas: {len(variables_encontradas)}")
    for var in sorted(variables_encontradas):
        print(f"   - {var}")
    
    return sorted(list(variables_encontradas))

# Probar con nuestro formato
if __name__ == "__main__":
    ruta_documento = "PactoMaker/formatos/Formato DESISTIMIENTO.docx"
    variables = encontrar_variables(ruta_documento)